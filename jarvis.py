import speech_recognition as sr
import os
import pyaudio
import webbrowser
import datetime
import openai
from config import apikey
import re
from docx import Document
import datetime
import time
from plyer import notification
from docx.shared import Pt
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
# import pywhatkit


chatStr = ""
# https://youtu.be/Z3ZAJoi4x6Q

def split_between_keywords(text, keyword1, keyword2):
    pattern = re.compile(f'{re.escape(keyword1)}(.*?){re.escape(keyword2)}', re.DOTALL)
    matches = pattern.findall(text)
    return matches


def takeCommand():
    r = sr.Recognizer()
    with sr.Microphone() as source:
        # r.pause_threshold =  0.6
        audio = r.listen(source)
        try:
            print("Recognizing...")
            print(audio)
            query = r.recognize_google(audio, language="en-in")
            print(f"User said: {query}")
            return query
        except Exception as e:
            return f"Some Error Occurred. Sorry from AI: {e}"

def chat(query):
    global chatStr
    print(chatStr)
    openai.api_key = apikey
    chatStr += f"Tatz: {query}\n Jarvis: "
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt= chatStr,
        temperature=0.7,
        max_tokens=256,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
    )
    # todo: Wrap this inside of a  try catch block
    try:
        chatresp = response['choices'][0]['text']
    except Exception as e:
        chatresp = e
    chatStr += f"{chatresp}\n"
    print (chatresp)
    return chatresp

def send_whatsapp_message(contact_name, message):
    print("test10")
    driver = webdriver.Chrome('/opt/homebrew/bin/chromedriver')
    print("test1")
    driver.get('https://web.whatsapp.com/')

    # Let user manually scan QR code
    input('Press Enter after scanning QR code...')

    # Find the contact
    search_box = WebDriverWait(driver, 50).until(EC.presence_of_element_located((By.XPATH, '//*[@id="side"]/div[1]/div/label/div/div[2]')))
    search_box.send_keys(contact_name)
    time.sleep(2)  # Give some time to load the search results

    # Click on the desired contact
    contact = driver.find_element(By.XPATH, f'//span[@title="{contact_name}"]')
    contact.click()

    # Find the input box and send the message
    msg_box = driver.find_element(By.XPATH, '//*[@id="main"]/footer/div[1]/div[2]/div/div[2]')
    msg_box.send_keys(message)
    msg_box.send_keys(Keys.ENTER)

    # Optionally, close the browser
    time.sleep(10)  # Wait to make sure message is sent
    driver.quit()

def formatString(a_string):
    escaped =  a_string.translate(str.maketrans({"-":  r"\-",
                                          "]":  r"\]",
                                          "\\": r"\\",
                                          "^":  r"\^",
                                          "$":  r"\$",
                                          "*":  r"\*",
                                          "'":  r"\'",
                                          ".":  r"\."}))
    return escaped.replace("\n", "")

def ai(prompt):
    openai.api_key = apikey
    text = f"OpenAI response for Prompt: {prompt} \n *************************\n\n"

    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=prompt,
        temperature=0.7,
        max_tokens=256,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
    )
    # todo: Wrap this inside of a  try catch block
    print(response["choices"][0]["text"])

    try:
        text += response["choices"][0]["text"]
    except Exception as e:
        text+=""
    if not os.path.exists("Openai"):
        os.mkdir("Openai")

    # with open(f"Openai/prompt- {random.randint(1, 2343434356)}", "w") as f:
    with open(f"Openai/{''.join(prompt.split('brain')[1:]).strip() }.txt", "w") as f:
        f.write(text)

    return response["choices"][0]["text"].lower()

def say(text):
    os.system(f"say -v Daniel {formatString(text)}")
    #os.system(f"say {formatString(text)}")

def set_reminder(reminder_text, minutes_from_now=1):
    
    due_date_script = ""
    
    if minutes_from_now > 0:
        # Let AppleScript calculate the due date.
        due_date_script = f"set the due date of newReminder to (current date + {minutes_from_now} * minutes)"

    # Using osascript to interact with AppleScript
    applescript_command = f"""
    tell application "Reminders"
        set newReminder to make new reminder with properties {{name:"{reminder_text}"}}
        {due_date_script}
    end tell
    """

    os.system(f"osascript -e '{applescript_command}'")

def set_alarm(hour, minute):
    """Set a calendar event that acts like an alarm."""
    applescript_command = f"""
    tell application "Calendar"
        tell calendar "MyCalendar"  # Replace with your calendar's name
            set theDate to current date
            set hours of theDate to {hour}
            set minutes of theDate to {minute}
            set seconds of theDate to 0
            make new event with properties {{summary:"Alarm", start date:theDate, end date:theDate + 60}}
        end tell
    end tell
    """

    os.system(f"osascript -e '{applescript_command}'")

def save_to_word(notes, filename='notes.docx'):
    doc = Document()
    doc.add_paragraph(notes)
    doc.save(filename)


def save_to_word_bullets(notes, filename='notes.docx'):
    doc = Document()
    
    # Separate the continuous string 'notes' into individual lines.
    lines = notes.strip().split('\n')
    
    # Add a bullet for each line.
    for line in lines:
        p = doc.add_paragraph()
        runner = p.add_run("• " + line)
        runner.font.size = Pt(12)
        
    doc.save(filename)




if __name__ == '__main__':
    print('Jarvis Activated')
    say("Hello, I am Jarvis. I am your personal assistant. How may I help you?")
    while True:
        print("Listening...")
        query = takeCommand()
        #query = "Jarvis send whatsapp"
        sites = [
            ["youtube", "https://www.youtube.com"], 
            ["wikipedia", "https://www.wikipedia.com"], 
            ["google", "https://www.google.com"],
            ["blockzen","https://blockzen.eu/"]
            ]
        if "jarvis".lower() in query.lower():
                for site in sites:
                    if f"Open {site[0]}".lower() in query.lower():
                        say (f"Opening {site[0]} Tatz")
                        webbrowser.open(site[1])
                        query = ""

                if "some music".lower() in query.lower():
                    musicPath = "/Users/tatz/Documents/rfm.mp3"
                    os.system(f"open {musicPath}")
                    query = ""
                
                elif "stop music".lower() in query.lower():
                    os.system(f"killall Music")
                    query = ""

                elif "the time".lower() in query.lower():
                    hour = datetime.datetime.now().strftime("%H")
                    min = datetime.datetime.now().strftime("%M")
                    say(f"Tatz the time is {hour} hour and {min} minutes")
                    query = ""

                elif "open notes".lower() in query.lower():
                    say("Opening Notes")
                    os.system(f"open /System/Applications/Notes.app")
                
                elif "close notes".lower() in query.lower():
                    say("Closing Notes")
                    os.system(f"killall Notes")
                
                # elif "whatsapp".lower() in query.lower():
                    #pywhatkit.sendwhatmsg("+919874767678", "Hello from Jarvis", 23, 50)
                    # contactName = split_between_keywords(query, " to ", " saying ")[0] #drinking water
                    # messag = query.split('saying ')[1].strip()
                    # print(contactName)
                    # print(messag)
                    # send_whatsapp_message(contactName, messag)
                    # query = ""
                
                elif "stop listening".lower() in query.lower():
                    say("I am going to sleep now and will be waiting for your command. See you soon. Good Byeee")
                    query = ""
                    exit()

                elif "brain".lower() in query.lower():
                    say(ai(prompt=query))
                    query = ""

                elif "facetime".lower() in query.lower():
                    say("Opening Facetime")
                    os.system(f"open /System/Applications/FaceTime.app")
                    query = ""
                
                elif "reset chat".lower() in query.lower():
                    chatStr = ""
                
                elif "set a reminder for".lower() in query.lower():
                    # Extracting reminder details. This is a naive way; consider improving it.
                    reminder_detail = query.split("set a reminder for")[1].strip()
                    
                    # Naive way of getting minutes, improve this based on your requirement.
                    if " at ".lower() in query.lower():
                        reminder_text = split_between_keywords(query, " for ", " at ")[0] #drinking water
                        quryTime = query.split('at ')[1:] #5 pm
                        amorpm = query[-4:]
                        hour = int(quryTime[0].split(':')[0]) if amorpm == "a.m." else int(quryTime[0].split(':')[0]) + 12
                        minute = int(quryTime[0].split(':')[1][:2])
                        now = datetime.datetime.now()
                        reminderTime = datetime.datetime(now.year, now.month, now.day, hour, minute)  # 17:00 is 5 PM
                        minutes_from_now = round((reminderTime - now).total_seconds()/60)
                    else:
                        minutes_match = re.search(r'in (\d+) minutes', reminder_detail)
                        minutes_from_now = int(minutes_match.group(1)) if minutes_match else 1
                        reminder_text = reminder_detail.split(" in")[0].strip() if minutes_match else reminder_detail

                    set_reminder(reminder_text, minutes_from_now)
                    say(f"Reminder set for: {reminder_text}")
                
                elif "set an alarm for".lower() in query.lower():
                # This is a naive way to extract hours and minutes. Consider refining it based on your requirements.
                    match = re.search(r"set an alarm for (\d+):(\d+)", query.lower())
                    if match:
                        hour, minute = map(int, match.groups())
                        amorpm = query[-4:]
                        # if amorpm == "p.m.":
                        #     hour += 12
                        set_alarm(hour, minute)
                        say(f"Alarm set for {hour}:{minute}")
                
                elif "bullets" in query.lower():
                    say("Starting dictation mode. Please start dictating. Say 'stop dictation' to finish.")
                    notes = ""
                    while True:
                        dictation = takeCommand()
                        if "stop dictation" in dictation.lower():
                            break
                        notes += dictation + "\n"
                    
                    # Here, you can also give a custom name or implement logic to provide unique filenames.
                    save_to_word_bullets(notes, filename='notes.docx')
                    say("Dictation saved to a Word document.")
                    query = ""
                
                elif "notes" in query.lower():
                    say("Starting dictation mode. Please start dictating. Say 'stop dictation' to finish.")
                    notes = ""
                    while True:
                        dictation = takeCommand()
                        if "stop dictation" in dictation.lower():
                            break
                        notes += dictation + "\n"
                    
                    # Here, you can also give a custom name or implement logic to provide unique filenames.
                    save_to_word(notes, filename='notes.docx')
                    say("Dictation saved to a Word document.")
                    query = ""

                else:
                    print("Chatting...")
                    say(chat(query))
                    query = ""

       